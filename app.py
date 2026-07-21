# -*- coding: utf-8 -*-
"""
每日督学小助手（可爱番茄风升级版）
功能亮点：
 1. 启动即进入「今日任务」页，只显示当天的学习任务
 2. 学习计划支持两种方式：导入计划文档（Day1/第1天分块）或手动逐日录入
 3. 智能进度结转：根据以往「计划量 - 实际量」自动增减当天计划
    （例：西班牙语 07-21 多学了 1 小时，则 07-22 的计划自动减少 1 小时）
 4. 番茄钟倒计时，结束/停止后一键把实际学习时长记入复盘
 5. 奶油粉色可爱界面：问候语、连续打卡天数、随机鼓励语
"""
import os
import json
import shutil
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime, date, timedelta
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from PIL import Image, ImageTk, ImageDraw, ImageFont
import matplotlib.font_manager as fm
import glob
import re
import random

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False
CHART_FONT = fm.FontProperties(family='Microsoft YaHei', size=12)
CHART_FONT_SM = fm.FontProperties(family='Microsoft YaHei', size=10)
CHART_FONT_TITLE = fm.FontProperties(family='Microsoft YaHei', size=16)

# ===== 奶油草莓色系 =====
PASTEL = {
    'bg':            '#FFF7F2',   # 奶油底
    'card_bg':       '#FFFFFF',
    'frame_bg':      '#FFEEF3',   # 浅粉
    'text':          '#8A6D6D',   # 软棕
    'title':         '#E87A90',   # 草莓红
    'subtitle':      '#F0A868',   # 蜜桃橙
    'accent':        '#F4B8C6',   # 粉
    'accent2':       '#FDE0E7',   # 浅粉2
    'button_bg':     '#FFE3EB',
    'button_active': '#FFD2DE',
    'button_fg':     '#C25E75',
    'entry_bg':      '#FFFBF8',
    'border':        '#F6D8DD',
    'highlight':     '#FCE4EA',
    'success':       '#9FD8B4',
    'success_bg':    '#E7F6EC',
    'warning':       '#F2C879',
    'chart_bg':      '#FFF7F2',
    'chart_colors':  ['#F4B8C6', '#F8D3A9', '#B5DFC0', '#AFD4E8', '#D8C6E8',
                      '#F8E3A1', '#F0B8A8', '#B8D8E0', '#E8C6D0', '#C8E0B8'],
    'header_bg':     '#FFE9EF',
    'delete_bg':     '#FBE7E3',
    'delete_active': '#F6D5CE',
    'delete_fg':     '#C08578',
}

FONT = 'Microsoft YaHei'
PAD = 25
DEFAULT_TASK_MINUTES = 30   # 分点任务默认时长（文档未标注时间时使用）

ENCOURAGE = [
    "今天也要元气满满哦！(๑•̀ㅂ•́)و✧",
    "积少成多，你超棒的！🌟",
    "坚持就是胜利，加油鸭！🐥",
    "每一个不曾起舞的日子，都是对生命的辜负~ 💃",
    "学完这一页，奖励自己一朵小红花 🌸",
    "你努力的样子，真的会发光 ✨",
    "慢慢来，比较快。稳步前进！🐢",
    "今天的汗水，是明天的光芒 ☀️",
]

def _bold(n, s=True):
    return (FONT, n, 'bold')
def _norm(n):
    return (FONT, n)

def _create_wallpaper_tile(size=256):
    """生成奶油粉可爱无缝壁纸：渐变 + 小爱心 + 小星星 + 圆点"""
    img = Image.new('RGB', (size, size), (255, 247, 242))
    draw = ImageDraw.Draw(img)
    random.seed(42)

    for y in range(size):
        r = 255
        g = 247 - int(6 * (y / size))
        b = 242 - int(10 * (y / size))
        draw.line([(0, y), (size, y)], fill=(r, g, b))

    # 小爱心
    for _ in range(6):
        cx = random.randint(12, size - 12)
        cy = random.randint(12, size - 12)
        s = random.randint(3, 5)
        c = random.choice([(250, 205, 215), (248, 190, 200), (245, 215, 200)])
        draw.ellipse([cx - s, cy - s, cx, cy], fill=c)
        draw.ellipse([cx, cy - s, cx + s, cy], fill=c)
        draw.polygon([(cx - s, cy - s // 3), (cx + s, cy - s // 3), (cx, cy + s)], fill=c)

    # 小星星
    for _ in range(10):
        x = random.randint(5, size - 5)
        y = random.randint(5, size - 5)
        cs = random.choice([(245, 215, 170), (250, 200, 210), (210, 225, 235), (215, 230, 205)])
        s = random.randint(2, 3)
        draw.line([(x - s - 1, y), (x + s + 1, y)], fill=cs, width=1)
        draw.line([(x, y - s - 1), (x, y + s + 1)], fill=cs, width=1)
        draw.ellipse([x - 1, y - 1, x + 1, y + 1], fill=cs)

    # 小圆点
    for _ in range(16):
        x = random.randint(3, size - 3)
        y = random.randint(3, size - 3)
        c = random.choice([(250, 220, 225), (245, 230, 210), (220, 235, 225), (230, 225, 240)])
        r = random.randint(1, 3)
        draw.ellipse([x - r, y - r, x + r, y + r], fill=c)

    return img


class StudyTrackerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("🍅 每日督学小助手")
        self.root.geometry("1240x860")
        self.root.minsize(1120, 720)

        self.data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "example")
        self.daily_dir = os.path.join(self.data_dir, "daily")
        self.projects_file = os.path.join(self.data_dir, "projects.json")
        self.settings_file = os.path.join(self.data_dir, "settings.json")
        self.load_settings()

        self.today = date.today().strftime("%Y-%m-%d")
        self.today_display = date.today().strftime("%Y年%m月%d日")
        self.projects = {}
        self.daily_plan = {}      # key -> {project, task, time, base, adjust, completed, source}
        self.review_data = {}
        self.all_review_data = {}
        self.timer_running = False
        self._timer_after_id = None
        self.timer_remaining = 0
        self.timer_total = 0
        self.timer_task_key = None
        self.selected_project = None
        self.adjust_notes = []    # 今日计划自动调整说明

        self._setup_wallpaper()
        self.setup_styles()
        self.create_widgets()
        self.load_projects()
        self.load_daily_data()
        # 启动即只展示当天任务，并按以往进度自动调整
        self.refresh_daily_plan(auto=True)
        self.notebook.select(self.today_frame)

    # ===== 基础 =====
    def _ensure_dirs(self):
        for d in [self.data_dir, self.daily_dir]:
            if not os.path.exists(d):
                os.makedirs(d)

    def load_settings(self):
        self._ensure_dirs()
        self.start_date = date.today().strftime("%Y-%m-%d")
        if os.path.exists(self.settings_file):
            try:
                with open(self.settings_file, 'r', encoding='utf-8') as f:
                    s = json.load(f)
                    self.data_dir = s.get('data_dir', self.data_dir)
                    self.daily_dir = os.path.join(self.data_dir, "daily")
                    self.projects_file = os.path.join(self.data_dir, "projects.json")
                    self.settings_file = os.path.join(self.data_dir, "settings.json")
                    self.start_date = s.get('start_date', self.start_date)
            except Exception:
                pass
        self._ensure_dirs()

    def save_settings(self):
        with open(self.settings_file, 'w', encoding='utf-8') as f:
            json.dump({'data_dir': self.data_dir, 'start_date': self.start_date},
                      f, ensure_ascii=False, indent=2)

    def _greeting(self):
        h = datetime.now().hour
        if h < 6:
            return "夜深啦", "🌙"
        if h < 9:
            return "早上好", "🌅"
        if h < 12:
            return "上午好", "🌞"
        if h < 14:
            return "中午好", "🍱"
        if h < 18:
            return "下午好", "☕"
        return "晚上好", "🌙"

    def _setup_wallpaper(self):
        tile = _create_wallpaper_tile(256)
        w = max(self.root.winfo_screenwidth(), 1280)
        h = max(self.root.winfo_screenheight(), 900)
        full = Image.new('RGB', (w, h))
        for x in range(0, w, 256):
            for y in range(0, h, 256):
                full.paste(tile, (x, y))
        self.wallpaper = ImageTk.PhotoImage(full)
        self.wallpaper_label = tk.Label(self.root, image=self.wallpaper, borderwidth=0)
        self.wallpaper_label.place(x=0, y=0, relwidth=1, relheight=1)

    def setup_styles(self):
        self.style = ttk.Style()
        self.style.theme_use('clam')

        self.style.configure('Cute.TButton',
            font=_norm(13), background=PASTEL['button_bg'], foreground=PASTEL['button_fg'],
            borderwidth=0, padding=(18, 9))
        self.style.map('Cute.TButton',
            background=[('active', PASTEL['button_active']), ('pressed', PASTEL['accent2'])],
            foreground=[('active', PASTEL['text'])])

        self.style.configure('Delete.TButton',
            font=_norm(13), background=PASTEL['delete_bg'], foreground=PASTEL['delete_fg'],
            borderwidth=0, padding=(18, 9))
        self.style.map('Delete.TButton', background=[('active', PASTEL['delete_active'])])

        self.style.configure('Accent.TButton',
            font=_bold(13), background=PASTEL['title'], foreground='white',
            borderwidth=0, padding=(18, 9))
        self.style.map('Accent.TButton',
            background=[('active', '#D76A80'), ('pressed', '#C25E75')])

        self.style.configure('Cute.TEntry', font=_norm(13), padding=8, fieldbackground=PASTEL['entry_bg'])
        self.style.configure('Cute.TCombobox', font=_norm(13), padding=8, fieldbackground=PASTEL['entry_bg'])
        self.style.configure('Cute.TFrame', background=PASTEL['bg'])
        self.style.configure('Cute.TRadiobutton',
            font=_norm(13), background=PASTEL['bg'], foreground=PASTEL['text'])
        self.style.configure('Cute.TNotebook', background=PASTEL['bg'], borderwidth=0)
        self.style.configure('Cute.TNotebook.Tab',
            font=_bold(13), padding=[22, 11],
            background=PASTEL['frame_bg'], foreground=PASTEL['text'], borderwidth=0)
        self.style.map('Cute.TNotebook.Tab',
            background=[('selected', PASTEL['bg']), ('active', PASTEL['card_bg'])],
            foreground=[('selected', PASTEL['title'])])

        self.style.configure('Treeview',
            font=_norm(12), rowheight=34,
            background=PASTEL['card_bg'], fieldbackground=PASTEL['card_bg'],
            foreground=PASTEL['text'])
        self.style.configure('Treeview.Heading',
            font=_bold(12), background=PASTEL['frame_bg'], foreground=PASTEL['title'])
        self.style.map('Treeview', background=[('selected', PASTEL['accent2'])])

        self.root.configure(bg=PASTEL['bg'])

    def _make_cute_icon(self, size):
        """用PIL画一只可爱的小番茄图标"""
        img = Image.new('RGBA', (size, size), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle([2, 2, size - 2, size - 2], radius=16, fill=PASTEL['frame_bg'])
        cx, cy, r = size // 2, size // 2 + 6, size // 2 - 14
        # 番茄身体
        draw.ellipse([cx - r, cy - r, cx + r, cy + r], fill='#F27E8D')
        # 高光
        draw.ellipse([cx - r + 6, cy - r + 6, cx - r + 16, cy - r + 18], fill='#F8A8B4')
        # 叶子
        draw.ellipse([cx - 4, cy - r - 8, cx + 4, cy - r + 2], fill='#8FCB9B')
        draw.polygon([(cx, cy - r - 12), (cx - 10, cy - r - 2), (cx + 10, cy - r - 2)], fill='#8FCB9B')
        # 眼睛
        draw.ellipse([cx - 10, cy - 4, cx - 4, cy + 2], fill='#5A3A3A')
        draw.ellipse([cx + 4, cy - 4, cx + 10, cy + 2], fill='#5A3A3A')
        # 腮红
        draw.ellipse([cx - 16, cy + 3, cx - 8, cy + 8], fill='#F4A0AC')
        draw.ellipse([cx + 8, cy + 3, cx + 16, cy + 8], fill='#F4A0AC')
        # 嘴
        draw.arc([cx - 5, cy + 2, cx + 5, cy + 10], start=0, end=180, fill='#5A3A3A', width=2)
        return ImageTk.PhotoImage(img)

    # ===== 界面骨架 =====
    def create_widgets(self):
        self.icon_photo = self._make_cute_icon(80)
        self._date_labels = []        # 各页头部的日期标签（随系统时间实时刷新）

        self.notebook = ttk.Notebook(self.root, style='Cute.TNotebook')
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=20, pady=(10, 20))

        self.today_frame = ttk.Frame(self.notebook, style='Cute.TFrame')
        self.plan_frame = ttk.Frame(self.notebook, style='Cute.TFrame')
        self.review_frame = ttk.Frame(self.notebook, style='Cute.TFrame')
        self.settings_frame = ttk.Frame(self.notebook, style='Cute.TFrame')

        self.notebook.add(self.today_frame, text='  🌸 今日任务  ')
        self.notebook.add(self.plan_frame, text='  📚 学习计划  ')
        self.notebook.add(self.review_frame, text='  📊 学习复盘  ')
        self.notebook.add(self.settings_frame, text='  ⚙️ 系统设置  ')

        # 各页内容放入可滚动容器，内容再多也能完整展示
        self.today_inner = self._make_scrollable(self.today_frame)
        self.plan_inner = self._make_scrollable(self.plan_frame)
        self.review_inner = self._make_scrollable(self.review_frame)

        self._today_tab()
        self._plan_tab()
        self._review_tab()
        self._settings_tab()
        self._clock_tick()

    def _make_scrollable(self, parent):
        """Canvas + 滚动条 的可滚动容器，返回内部 Frame"""
        canvas = tk.Canvas(parent, bg=PASTEL['bg'], highlightthickness=0, bd=0)
        vsb = ttk.Scrollbar(parent, orient='vertical', command=canvas.yview)
        inner = tk.Frame(canvas, bg=PASTEL['bg'])
        win_id = canvas.create_window((0, 0), window=inner, anchor='nw')
        inner.bind('<Configure>',
                   lambda e: canvas.configure(scrollregion=canvas.bbox('all')))
        canvas.bind('<Configure>',
                    lambda e: canvas.itemconfig(win_id, width=e.width))
        canvas.configure(yscrollcommand=vsb.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)

        def _wheel(e):
            canvas.yview_scroll(int(-1 * (e.delta / 120)), 'units')
        canvas.bind('<Enter>', lambda e: canvas.bind_all('<MouseWheel>', _wheel))
        canvas.bind('<Leave>', lambda e: canvas.unbind_all('<MouseWheel>'))
        return inner

    def _clock_tick(self):
        """每30秒与电脑时间对时：刷新问候语/日期；跨天自动切换到新的一天"""
        now = date.today()
        if now.strftime("%Y-%m-%d") != self.today:
            self.today = now.strftime("%Y-%m-%d")
            self.today_display = now.strftime("%Y年%m月%d日")
            self.daily_plan = {}
            self.review_data = {}
            self.load_daily_data()
            self.refresh_daily_plan(auto=True)
        greet, emoji = self._greeting()
        if hasattr(self, 'today_title_label'):
            self.today_title_label.config(text=f'{emoji} {greet}，今天也要加油鸭！')
        for lb in self._date_labels:
            lb.config(text=f'📅 {self.today_display}')
        self.root.after(30000, self._clock_tick)

    def _header(self, parent, title, icon=True, right_text=None):
        h = tk.Frame(parent, bg=PASTEL['header_bg'], highlightthickness=0)
        h.pack(fill=tk.X, padx=PAD, pady=(10, 4))
        if icon:
            tk.Label(h, image=self.icon_photo, bg=PASTEL['header_bg']).pack(side=tk.LEFT, padx=10, pady=10)
        title_lb = tk.Label(h, text=title, font=_bold(24), fg=PASTEL['title'],
                 bg=PASTEL['header_bg'])
        title_lb.pack(side=tk.LEFT, padx=15, pady=10)
        date_lb = tk.Label(h, text=right_text if right_text is not None else f'📅 {self.today_display}',
                 font=_bold(15), fg=PASTEL['subtitle'], bg=PASTEL['header_bg'])
        date_lb.pack(side=tk.RIGHT, padx=20, pady=10)
        if right_text is None or right_text.startswith('📅'):
            self._date_labels.append(date_lb)
        return title_lb

    def _labelframe(self, parent, text, **kw):
        return tk.LabelFrame(parent, text=text, font=_bold(13), fg=PASTEL['title'],
            bg=PASTEL['card_bg'], padx=20, pady=10, relief='solid', bd=1, **kw)

    def _entry(self, parent, var, width=20):
        return tk.Entry(parent, textvariable=var, width=width, font=_norm(13),
            bg=PASTEL['entry_bg'], relief='solid', bd=1, fg=PASTEL['text'],
            insertbackground=PASTEL['text'])

    def _label(self, parent, text, **kw):
        return tk.Label(parent, text=text, font=_norm(13), fg=PASTEL['text'],
            bg=PASTEL['card_bg'], **kw)

    # ===== 今日任务页（默认首页） =====
    def _today_tab(self):
        greet, emoji = self._greeting()
        self.today_title_label = self._header(
            self.today_inner, f'{emoji} {greet}，今天也要加油鸭！', icon=True)

        # 顶部统计卡片
        stat = tk.Frame(self.today_inner, bg=PASTEL['bg'])
        stat.pack(fill=tk.X, padx=PAD, pady=3)
        self.stat_cards = {}
        for key, txt, color in [
            ('total', '今日任务', PASTEL['accent2']),
            ('done', '已完成', PASTEL['success_bg']),
            ('mins', '计划总时长', '#FDF3E0'),
            ('streak', '连续学习', '#F0EBFA'),
        ]:
            card = tk.Frame(stat, bg=color, highlightbackground=PASTEL['border'],
                            highlightthickness=1)
            card.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=6, ipady=2)
            num = tk.Label(card, text='0', font=_bold(19), fg=PASTEL['title'], bg=color)
            num.pack(pady=(2, 0))
            tk.Label(card, text=txt, font=_norm(11), fg=PASTEL['text'], bg=color).pack(pady=(0, 2))
            self.stat_cards[key] = num

        # 自动调整提示条
        self.adjust_var = tk.StringVar(value='')
        self.adjust_label = tk.Label(self.today_inner, textvariable=self.adjust_var,
            font=_norm(11), fg=PASTEL['subtitle'], bg=PASTEL['bg'], justify=tk.LEFT)
        self.adjust_label.pack(fill=tk.X, padx=PAD, pady=(2, 0))

        # 今日任务列表
        lf = self._labelframe(self.today_inner,
            '📋 今日学习任务（点击「▶」开始计时，点击其他任务自动切换；点击「打卡」完成）')
        lf.pack(fill=tk.BOTH, expand=True, padx=PAD, pady=4)

        cols = ('name', 'base', 'adjust', 'time', 'status', 'timer', 'action')
        self.checkin_tree = ttk.Treeview(lf, columns=cols, show='headings', height=6)
        self.checkin_tree.heading('name', text='🌷 学习任务')
        self.checkin_tree.heading('base', text='原计划')
        self.checkin_tree.heading('adjust', text='智能调整')
        self.checkin_tree.heading('time', text='今日计划')
        self.checkin_tree.heading('status', text='状态')
        self.checkin_tree.heading('timer', text='计时')
        self.checkin_tree.heading('action', text='操作')
        self.checkin_tree.column('name', width=300)
        self.checkin_tree.column('base', width=90, anchor='center')
        self.checkin_tree.column('adjust', width=100, anchor='center')
        self.checkin_tree.column('time', width=100, anchor='center')
        self.checkin_tree.column('status', width=95, anchor='center')
        self.checkin_tree.column('timer', width=95, anchor='center')
        self.checkin_tree.column('action', width=90, anchor='center')
        sb = ttk.Scrollbar(lf, orient='vertical', command=self.checkin_tree.yview)
        self.checkin_tree.configure(yscrollcommand=sb.set)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.checkin_tree.pack(fill=tk.BOTH, expand=True)
        self.checkin_tree.tag_configure('done', background=PASTEL['success_bg'])
        self.checkin_tree.tag_configure('timing', background='#FFF1D6')
        self.checkin_tree.bind('<ButtonRelease-1>', self._on_checkin_click)

        # 手动添加今日临时任务
        form = tk.Frame(self.today_inner, bg=PASTEL['bg'])
        form.pack(fill=tk.X, padx=PAD, pady=2)
        self.checkin_name_var = tk.StringVar()
        self.checkin_time_var = tk.StringVar()
        tk.Label(form, text='➕ 临时任务', font=_norm(13), fg=PASTEL['text'],
                 bg=PASTEL['bg']).pack(side=tk.LEFT, padx=(0, 8))
        self._entry(form, self.checkin_name_var, 26).pack(side=tk.LEFT, padx=4)
        tk.Label(form, text='分钟', font=_norm(13), fg=PASTEL['text'], bg=PASTEL['bg']).pack(side=tk.LEFT, padx=4)
        self._entry(form, self.checkin_time_var, 8).pack(side=tk.LEFT, padx=4)
        ttk.Button(form, text='添加', command=self.add_daily_item,
                   style='Cute.TButton').pack(side=tk.LEFT, padx=8)
        ttk.Button(form, text='删除选中', command=self.delete_daily_item,
                   style='Delete.TButton').pack(side=tk.LEFT, padx=4)
        ttk.Button(form, text='🔄 重新同步计划', command=lambda: self.refresh_daily_plan(auto=False),
                   style='Cute.TButton').pack(side=tk.RIGHT, padx=4)

        # 番茄钟
        tf = tk.LabelFrame(self.today_inner,
            text=' 🍅 番茄学习钟（计时结束自动打卡并记录；切换任务自动记录上一段） ',
            font=_bold(12), fg=PASTEL['title'], bg=PASTEL['card_bg'],
            padx=20, pady=8, relief='solid', bd=1)
        tf.pack(fill=tk.X, padx=PAD, pady=(2, 8))

        row = tk.Frame(tf, bg=PASTEL['card_bg'])
        row.pack(fill=tk.X)
        self.timer_display = tk.Label(row, text='00:00:00',
            font=('Arial', 34, 'bold'), fg=PASTEL['title'], bg=PASTEL['card_bg'])
        self.timer_display.pack(side=tk.LEFT, padx=(10, 25))

        tc = tk.Frame(row, bg=PASTEL['card_bg'])
        tc.pack(side=tk.LEFT, pady=4)
        self.checkin_status_var = tk.StringVar(value=random.choice(ENCOURAGE))
        tk.Label(tc, textvariable=self.checkin_status_var, font=_norm(12),
                 fg=PASTEL['subtitle'], bg=PASTEL['card_bg']).pack(anchor=tk.W, pady=(0, 4))
        btns = tk.Frame(tc, bg=PASTEL['card_bg'])
        btns.pack(anchor=tk.W)
        self.pause_timer_btn = ttk.Button(btns, text='⏸ 暂停', command=self.pause_timer,
                                          style='Cute.TButton', state=tk.DISABLED)
        self.pause_timer_btn.pack(side=tk.LEFT, padx=6)
        self.stop_timer_btn = ttk.Button(btns, text='⏹ 结束并记录', command=self.stop_timer,
                                         style='Delete.TButton', state=tk.DISABLED)
        self.stop_timer_btn.pack(side=tk.LEFT, padx=6)

    # ===== 学习计划页 =====
    def _plan_tab(self):
        self._header(self.plan_inner, '📚 学习计划', icon=False)

        top = tk.Frame(self.plan_inner, bg=PASTEL['bg'])
        top.pack(fill=tk.BOTH, expand=True, padx=PAD, pady=5)

        # 左栏：项目管理
        left = tk.Frame(top, bg=PASTEL['bg'])
        left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))

        pf = self._labelframe(left, '🌱 学习项目管理')
        pf.pack(fill=tk.X, pady=5)
        self._label(pf, '项目名称').grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        self.proj_name_var = tk.StringVar()
        self._entry(pf, self.proj_name_var, 22).grid(row=0, column=1, padx=10, pady=10)
        bf = tk.Frame(pf, bg=PASTEL['card_bg'])
        bf.grid(row=1, column=0, columnspan=2, pady=8)
        ttk.Button(bf, text='添加项目', command=self.add_project,
                   style='Cute.TButton').pack(side=tk.LEFT, padx=8)
        ttk.Button(bf, text='删除项目', command=self.remove_project,
                   style='Delete.TButton').pack(side=tk.LEFT, padx=8)

        self.proj_listbox = tk.Listbox(left, font=_norm(13), bg=PASTEL['card_bg'],
            fg=PASTEL['text'], selectbackground=PASTEL['accent2'], selectforeground=PASTEL['text'],
            activestyle='none', relief='solid', bd=1)
        plb_sb = ttk.Scrollbar(left, orient='vertical', command=self.proj_listbox.yview)
        self.proj_listbox.configure(yscrollcommand=plb_sb.set)
        plb_sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.proj_listbox.pack(fill=tk.BOTH, expand=True, pady=5)
        self.proj_listbox.bind('<<ListboxSelect>>', self._on_proj_select)

        # 右栏：项目计划（文档导入 + 手动录入）
        right = tk.Frame(top, bg=PASTEL['bg'])
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(10, 0))

        self.plan_info_label = tk.Label(right, text='👈 请先选择或创建一个学习项目',
            font=_bold(14), fg=PASTEL['accent'], bg=PASTEL['bg'])
        self.plan_info_label.pack(anchor=tk.W, pady=5)

        # --- 方式一：导入文档 ---
        df = self._labelframe(right, '📄 方式一：导入计划文档')
        df.pack(fill=tk.BOTH, expand=True, pady=5)
        row1 = tk.Frame(df, bg=PASTEL['card_bg'])
        row1.pack(fill=tk.X, pady=4)
        self.import_btn = ttk.Button(row1, text='导入计划文件 (.txt/.docx)',
            command=self.import_plan_for_project, style='Cute.TButton', state=tk.DISABLED)
        self.import_btn.pack(side=tk.LEFT, padx=5)
        self.del_doc_btn = ttk.Button(row1, text='删除文档', command=self.delete_project_doc,
            style='Delete.TButton', state=tk.DISABLED)
        self.del_doc_btn.pack(side=tk.LEFT, padx=5)
        self.start_date_btn = ttk.Button(row1, text='📅 调整起始日期',
            command=self.edit_project_start_date, style='Cute.TButton', state=tk.DISABLED)
        self.start_date_btn.pack(side=tk.LEFT, padx=5)

        self.doc_listbox = tk.Listbox(df, font=_norm(12), bg=PASTEL['card_bg'],
            fg=PASTEL['text'], selectbackground=PASTEL['accent2'], selectforeground=PASTEL['text'],
            activestyle='none', relief='solid', bd=1, height=4)
        dlb_sb = ttk.Scrollbar(df, orient='vertical', command=self.doc_listbox.yview)
        self.doc_listbox.configure(yscrollcommand=dlb_sb.set)
        dlb_sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.doc_listbox.pack(fill=tk.BOTH, expand=True, pady=4)
        tk.Label(df, text='支持格式：① [ ] Day 1：任务1｜任务2（兼容 - [ ] **Day 1**、Day 1-2 区间）　② Day1 单独一行 + 每行「任务名,分钟数」　③ 日期排课：7.25 瑜伽 16:00-17:30',
            font=_norm(11), fg=PASTEL['subtitle'], bg=PASTEL['card_bg']).pack(anchor=tk.W, pady=2)
        tk.Label(df, text='💡 重新上传即更新为最新版本；Day按项目起始日期推算，中断了可点「调整起始日期」续上',
            font=_norm(11), fg=PASTEL['subtitle'], bg=PASTEL['card_bg']).pack(anchor=tk.W, pady=2)

        row0 = tk.Frame(df, bg=PASTEL['card_bg'])
        row0.pack(fill=tk.X, pady=2)
        tk.Label(row0, text='⏱️ 分点任务默认时长(分钟)（文档未标注"限时X分钟/X小时"时使用）',
            font=_norm(11), fg=PASTEL['text'], bg=PASTEL['card_bg']).pack(side=tk.LEFT)
        self.default_min_var = tk.StringVar(value=str(DEFAULT_TASK_MINUTES))
        self._entry(row0, self.default_min_var, 6).pack(side=tk.LEFT, padx=6)
        self.default_min_btn = ttk.Button(row0, text='保存', command=self.save_default_minutes,
            style='Cute.TButton', state=tk.DISABLED)
        self.default_min_btn.pack(side=tk.LEFT)

        # --- 方式二：手动逐日录入 ---
        mf = self._labelframe(right, '✏️ 方式二：手动录入每日计划')
        mf.pack(fill=tk.X, pady=5)
        row2 = tk.Frame(mf, bg=PASTEL['card_bg'])
        row2.pack(fill=tk.X, pady=4)
        tk.Label(row2, text='第', font=_norm(13), fg=PASTEL['text'],
                 bg=PASTEL['card_bg']).pack(side=tk.LEFT)
        self.manual_day_var = tk.StringVar(value='1')
        self._entry(row2, self.manual_day_var, 5).pack(side=tk.LEFT, padx=4)
        tk.Label(row2, text='天', font=_norm(13), fg=PASTEL['text'],
                 bg=PASTEL['card_bg']).pack(side=tk.LEFT)
        self.manual_task_var = tk.StringVar()
        self._entry(row2, self.manual_task_var, 18).pack(side=tk.LEFT, padx=6)
        tk.Label(row2, text='分钟', font=_norm(13), fg=PASTEL['text'],
                 bg=PASTEL['card_bg']).pack(side=tk.LEFT)
        self.manual_time_var = tk.StringVar()
        self._entry(row2, self.manual_time_var, 7).pack(side=tk.LEFT, padx=4)
        self.manual_add_btn = ttk.Button(row2, text='加入计划', command=self.add_manual_plan_item,
            style='Cute.TButton', state=tk.DISABLED)
        self.manual_add_btn.pack(side=tk.LEFT, padx=6)
        self.manual_del_btn = ttk.Button(row2, text='删除选中', command=self.delete_manual_plan_item,
            style='Delete.TButton', state=tk.DISABLED)
        self.manual_del_btn.pack(side=tk.LEFT, padx=2)

        self.manual_listbox = tk.Listbox(mf, font=_norm(12), bg=PASTEL['card_bg'],
            fg=PASTEL['text'], selectbackground=PASTEL['accent2'], selectforeground=PASTEL['text'],
            activestyle='none', relief='solid', bd=1, height=6)
        mlb_sb = ttk.Scrollbar(mf, orient='vertical', command=self.manual_listbox.yview)
        self.manual_listbox.configure(yscrollcommand=mlb_sb.set)
        mlb_sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.manual_listbox.pack(fill=tk.BOTH, expand=True, pady=4)
        tk.Label(mf, text='💡 手动计划与文档计划会合并；当天实际与计划的差额会自动结转到之后的计划里',
            font=_norm(11), fg=PASTEL['subtitle'], bg=PASTEL['card_bg']).pack(anchor=tk.W, pady=2)

    # ===== 学习复盘页 =====
    def _review_tab(self):
        self._header(self.review_inner, '📊 学习复盘', icon=False)

        ch = tk.Frame(self.review_inner, bg=PASTEL['bg'])
        ch.pack(fill=tk.X, padx=PAD, pady=5)
        self.chart_mode_var = tk.StringVar(value='today')
        ttk.Radiobutton(ch, text='⏰ 今日24小时分布', variable=self.chart_mode_var, value='today',
            command=self.update_review_chart, style='Cute.TRadiobutton').pack(side=tk.LEFT, padx=12)
        ttk.Radiobutton(ch, text='📈 整体进度', variable=self.chart_mode_var, value='all',
            command=self.update_review_chart, style='Cute.TRadiobutton').pack(side=tk.LEFT, padx=12)

        form = self._labelframe(self.review_inner, '🖊️ 复盘记录（选中下方记录可修改/删除）')
        form.pack(fill=tk.X, padx=PAD, pady=5)
        self._label(form, '学习项目').grid(row=0, column=0, padx=10, pady=8, sticky=tk.W)
        self.review_name_var = tk.StringVar()
        self.review_name_combo = ttk.Combobox(form, textvariable=self.review_name_var,
            width=25, style='Cute.TCombobox')
        self.review_name_combo.grid(row=0, column=1, padx=10, pady=8)
        self._label(form, '实际时间(分钟)').grid(row=0, column=2, padx=10, pady=8, sticky=tk.W)
        self.review_time_var = tk.StringVar()
        self._entry(form, self.review_time_var, 10).grid(row=0, column=3, padx=10, pady=8)
        self._label(form, '学习内容').grid(row=1, column=0, padx=10, pady=8, sticky=tk.W)
        self.review_content_var = tk.StringVar()
        self._entry(form, self.review_content_var, 48).grid(row=1, column=1, padx=10, pady=8, columnspan=3)
        self.add_review_btn = ttk.Button(form, text='添加复盘', command=self.add_review_item,
                   style='Cute.TButton')
        self.add_review_btn.grid(row=1, column=4, padx=15, pady=8)
        self._editing_review = None   # (任务名, 记录索引)，None 表示新增模式

        tree_box = tk.Frame(self.review_inner, bg=PASTEL['bg'])
        tree_box.pack(fill=tk.X, padx=PAD, pady=5)
        self.review_tree = ttk.Treeview(tree_box, columns=('name', 'time', 'content'),
            show='headings', height=5)
        self.review_tree.heading('name', text='学习项目')
        self.review_tree.heading('time', text='实际时间 (分钟)')
        self.review_tree.heading('content', text='学习内容')
        self.review_tree.column('name', width=220)
        self.review_tree.column('time', width=130, anchor='center')
        self.review_tree.column('content', width=400)
        rsb = ttk.Scrollbar(tree_box, orient='vertical', command=self.review_tree.yview)
        self.review_tree.configure(yscrollcommand=rsb.set)
        rsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.review_tree.pack(fill=tk.X, expand=True)
        self.review_tree.bind('<Double-1>', lambda e: self.edit_review_item())

        rb = tk.Frame(self.review_inner, bg=PASTEL['bg'])
        rb.pack(fill=tk.X, padx=PAD, pady=(0, 5))
        tk.Label(rb, text='💡 双击记录可直接修改', font=_norm(11),
                 fg=PASTEL['subtitle'], bg=PASTEL['bg']).pack(side=tk.LEFT)
        ttk.Button(rb, text='✏️ 修改选中', command=self.edit_review_item,
                   style='Cute.TButton').pack(side=tk.RIGHT, padx=4)
        ttk.Button(rb, text='🗑️ 删除选中', command=self.delete_review_item,
                   style='Delete.TButton').pack(side=tk.RIGHT, padx=4)

        # 今日项目统计：汇总该项目全部任务时间 → 自动结转后续计划
        sf = self._labelframe(self.review_inner,
            '🧮 今日项目统计（实际与计划的差额会自动调整之后的打卡时间）')
        sf.pack(fill=tk.X, padx=PAD, pady=5)
        self.summary_tree = ttk.Treeview(sf,
            columns=('proj', 'planned', 'actual', 'diff', 'effect'),
            show='headings', height=3)
        self.summary_tree.heading('proj', text='学习项目')
        self.summary_tree.heading('planned', text='今日计划')
        self.summary_tree.heading('actual', text='今日实际')
        self.summary_tree.heading('diff', text='差额')
        self.summary_tree.heading('effect', text='对后续计划的影响')
        self.summary_tree.column('proj', width=220)
        self.summary_tree.column('planned', width=100, anchor='center')
        self.summary_tree.column('actual', width=100, anchor='center')
        self.summary_tree.column('diff', width=100, anchor='center')
        self.summary_tree.column('effect', width=330)
        self.summary_tree.pack(fill=tk.X)

        cf = self._labelframe(self.review_inner, '🍭 统计图表')
        cf.pack(fill=tk.BOTH, expand=True, padx=PAD, pady=5, ipady=5)
        self.figure, self.ax = plt.subplots(figsize=(9, 5.2))
        self.figure.patch.set_facecolor(PASTEL['chart_bg'])
        self.canvas = FigureCanvasTkAgg(self.figure, cf)
        self.canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    # ===== 设置页 =====
    def _settings_tab(self):
        self._header(self.settings_frame, '⚙️ 系统设置', icon=False)
        sf = self._labelframe(self.settings_frame, '🗂️ 数据存储设置')
        sf.pack(fill=tk.X, padx=PAD, pady=20)
        self._label(sf, '数据文件夹').grid(row=0, column=0, padx=10, pady=14, sticky=tk.W)
        self.data_dir_var = tk.StringVar(value=self.data_dir)
        self._entry(sf, self.data_dir_var, 55).grid(row=0, column=1, padx=10, pady=14)
        ttk.Button(sf, text='浏览', command=self.browse_data_dir,
                   style='Cute.TButton').grid(row=0, column=2, padx=10, pady=14)
        ttk.Button(sf, text='默认路径', command=self.set_default_dir,
                   style='Cute.TButton').grid(row=0, column=3, padx=5, pady=14)
        ttk.Button(self.settings_frame, text='保存设置', command=self.save_settings_and_restart,
                   style='Accent.TButton').pack(side=tk.BOTTOM, pady=25)
        tk.Label(self.settings_frame, text='修改数据文件夹后需要重新启动应用生效',
            font=_norm(12), fg=PASTEL['accent'], bg=PASTEL['bg']).pack(side=tk.BOTTOM, pady=15)

    # ===== 项目管理 =====
    def load_projects(self):
        if os.path.exists(self.projects_file):
            try:
                with open(self.projects_file, 'r', encoding='utf-8') as f:
                    self.projects = json.load(f)
            except Exception:
                self.projects = {}
        else:
            self.projects = {}
        # 兼容旧数据，补充字段
        for pn, p in self.projects.items():
            p.setdefault('manual_schedule', {})
            p.setdefault('plans', [])
            p.setdefault('task_minutes', DEFAULT_TASK_MINUTES)
        self._ensure_project_dirs()
        self._update_proj_list()

    def save_projects(self):
        with open(self.projects_file, 'w', encoding='utf-8') as f:
            json.dump(self.projects, f, ensure_ascii=False, indent=2)

    def _ensure_project_dirs(self):
        for pn in self.projects:
            os.makedirs(self._proj_dir(pn), exist_ok=True)

    def _proj_dir(self, pn):
        d = os.path.join(self.data_dir, pn)
        os.makedirs(d, exist_ok=True)
        return d

    def _plan_files(self, proj_name):
        """项目计划文档列表，按上传时间（mtime）从新到旧排序"""
        files = glob.glob(os.path.join(self._proj_dir(proj_name), "plan_*"))
        return sorted(files, key=lambda f: os.path.getmtime(f)
                      if os.path.exists(f) else 0, reverse=True)

    def _calc_project_day(self, proj_name):
        """每个项目独立计算今天是第几天"""
        proj = self.projects.get(proj_name, {})
        sd = proj.get('start_date')
        if not sd:
            sd = date.today().strftime("%Y-%m-%d")
            if proj_name in self.projects:
                self.projects[proj_name]['start_date'] = sd
                self.save_projects()
        try:
            sd_date = datetime.strptime(sd, "%Y-%m-%d").date()
            return max((date.today() - sd_date).days + 1, 1)
        except Exception:
            return 1

    def add_project(self):
        name = self.proj_name_var.get().strip()
        if not name:
            messagebox.showwarning('提示', '请输入项目名称！')
            return
        if name in self.projects:
            messagebox.showwarning('提示', '该项目已存在！')
            return
        self.projects[name] = {
            'plans': [],
            'manual_schedule': {},
            'task_minutes': DEFAULT_TASK_MINUTES,
            'start_date': date.today().strftime("%Y-%m-%d"),
        }
        self._proj_dir(name)
        self.save_projects()
        self._update_proj_list()
        self.proj_name_var.set('')
        messagebox.showinfo('成功', f'项目「{name}」已创建！🌱')

    def remove_project(self):
        sel = self.proj_listbox.curselection()
        if not sel:
            messagebox.showwarning('提示', '请先选择要删除的项目！')
            return
        name = self.proj_listbox.get(sel[0]).split('  (')[0]
        if messagebox.askyesno('删除确认', f'确定要删除项目「{name}」及其所有计划文件吗？'):
            shutil.rmtree(self._proj_dir(name), ignore_errors=True)
            del self.projects[name]
            self.save_projects()
            self.selected_project = None
            self._update_proj_list()
            self._clear_plan_view()

    def _update_proj_list(self):
        self.proj_listbox.delete(0, tk.END)
        for name in sorted(self.projects.keys()):
            docs = glob.glob(os.path.join(self._proj_dir(name), "plan_*"))
            n_manual = sum(len(v) for v in self.projects[name].get('manual_schedule', {}).values())
            self.proj_listbox.insert(tk.END, f"{name}  ({len(docs)}文档/{n_manual}手动)")

    def _on_proj_select(self, event):
        sel = self.proj_listbox.curselection()
        if not sel:
            return
        self.selected_project = self.proj_listbox.get(sel[0]).split('  (')[0]
        for b in (self.import_btn, self.del_doc_btn, self.manual_add_btn,
                  self.manual_del_btn, self.default_min_btn, self.start_date_btn):
            b.config(state=tk.NORMAL)
        day = self._calc_project_day(self.selected_project)
        self.plan_info_label.config(
            text=f'已选择：{self.selected_project}　🌞 今天是第 {day} 天')
        self.default_min_var.set(str(self.projects[self.selected_project].get(
            'task_minutes', DEFAULT_TASK_MINUTES)))
        self.manual_day_var.set(str(day))
        self._update_doc_list()
        self._update_manual_list()

    def _update_doc_list(self):
        self.doc_listbox.delete(0, tk.END)
        if not self.selected_project:
            return
        plans = self._plan_files(self.selected_project)
        for p in plans:
            self.doc_listbox.insert(tk.END, os.path.basename(p))
        if not plans:
            self.doc_listbox.insert(tk.END, '（尚未导入计划文件）')

    def _update_manual_list(self):
        self.manual_listbox.delete(0, tk.END)
        if not self.selected_project:
            return
        sched = self.projects[self.selected_project].get('manual_schedule', {})
        for day in sorted(sched.keys(), key=lambda x: int(x)):
            for item in sched[day]:
                self.manual_listbox.insert(
                    tk.END, f"第{day}天 ｜ {item['name']} ｜ {item['time']}分钟")
        if not sched:
            self.manual_listbox.insert(tk.END, '（暂无手动计划，可在上方录入）')

    def save_default_minutes(self):
        if not self.selected_project:
            return
        v = self.default_min_var.get().strip()
        if not v.isdigit() or int(v) < 1:
            messagebox.showwarning('提示', '默认时长请输入正整数（分钟）！')
            return
        self.projects[self.selected_project]['task_minutes'] = int(v)
        self.save_projects()
        self.refresh_daily_plan(auto=True)
        messagebox.showinfo('已保存', f'「{self.selected_project}」分点任务默认时长已设为 {v} 分钟！')

    def edit_project_start_date(self):
        """调整项目起始日期：中断后重新上传计划时，用它对齐Day编号"""
        if not self.selected_project:
            return
        proj = self.projects[self.selected_project]
        win = tk.Toplevel(self.root)
        win.title(f'{self.selected_project} - 调整起始日期')
        win.geometry('460x260')
        win.configure(bg=PASTEL['bg'])
        win.transient(self.root)
        win.grab_set()
        tk.Label(win, text='📅 调整项目起始日期', font=_bold(16),
                 fg=PASTEL['title'], bg=PASTEL['bg']).pack(pady=(18, 6))
        tk.Label(win,
            text='Day1 = 起始日期当天，之后每过一天 Day+1。\n'
                 '学习中断了几天/几个月后，重新上传计划并把起始日期\n'
                 '往前调相应天数，进度就能无缝续上。',
            font=_norm(12), fg=PASTEL['text'], bg=PASTEL['bg'], justify=tk.CENTER).pack(pady=6)
        row = tk.Frame(win, bg=PASTEL['bg'])
        row.pack(pady=8)
        tk.Label(row, text='起始日期(YYYY-MM-DD)', font=_norm(12),
                 fg=PASTEL['text'], bg=PASTEL['bg']).pack(side=tk.LEFT, padx=6)
        var = tk.StringVar(value=proj.get('start_date', self.today))
        self._entry(row, var, 14).pack(side=tk.LEFT)

        def do_save():
            v = var.get().strip()
            try:
                datetime.strptime(v, "%Y-%m-%d")
            except Exception:
                messagebox.showwarning('提示', '日期格式应为 YYYY-MM-DD，如 2026-07-20')
                return
            proj['start_date'] = v
            self.save_projects()
            win.destroy()
            day = self._calc_project_day(self.selected_project)
            self.plan_info_label.config(
                text=f'已选择：{self.selected_project}　🌞 今天是第 {day} 天')
            self.manual_day_var.set(str(day))
            self.refresh_daily_plan(auto=True)
            messagebox.showinfo('已保存',
                f'起始日期已调整为 {v}，今天是第 {day} 天，\n今日任务已按最新计划同步！')

        ttk.Button(win, text='保存', command=do_save,
                   style='Accent.TButton').pack(side=tk.LEFT, padx=60, pady=14)
        ttk.Button(win, text='取消', command=win.destroy,
                   style='Cute.TButton').pack(side=tk.RIGHT, padx=60, pady=14)

    def _clear_plan_view(self):
        self.plan_info_label.config(text='👈 请先选择或创建一个学习项目')
        for b in (self.import_btn, self.del_doc_btn, self.manual_add_btn,
                  self.manual_del_btn, self.default_min_btn, self.start_date_btn):
            b.config(state=tk.DISABLED)
        self.doc_listbox.delete(0, tk.END)
        self.manual_listbox.delete(0, tk.END)

    # ===== 手动逐日计划 =====
    def add_manual_plan_item(self):
        if not self.selected_project:
            return
        day = self.manual_day_var.get().strip()
        task = self.manual_task_var.get().strip()
        tstr = self.manual_time_var.get().strip()
        if not day.isdigit() or int(day) < 1:
            messagebox.showwarning('提示', '天数请输入正整数！')
            return
        if not task:
            messagebox.showwarning('提示', '请输入任务名称！')
            return
        if not tstr.isdigit():
            messagebox.showwarning('提示', '时间请输入数字（分钟）！')
            return
        sched = self.projects[self.selected_project].setdefault('manual_schedule', {})
        sched.setdefault(day, []).append({'name': task, 'time': int(tstr)})
        self.save_projects()
        self._update_manual_list()
        self._update_proj_list()
        self.manual_task_var.set('')
        self.manual_time_var.set('')
        self.refresh_daily_plan(auto=True)
        messagebox.showinfo('已加入', f'已加入「{self.selected_project}」第{day}天计划！🌸')

    def delete_manual_plan_item(self):
        sel = self.manual_listbox.curselection()
        if not sel or not self.selected_project:
            messagebox.showwarning('提示', '请先选择要删除的计划项！')
            return
        line = self.manual_listbox.get(sel[0])
        m = re.match(r'第(\d+)天 ｜ (.+) ｜ (\d+)分钟', line)
        if not m:
            return
        day, name = m.group(1), m.group(2)
        sched = self.projects[self.selected_project].get('manual_schedule', {})
        if day in sched:
            sched[day] = [i for i in sched[day] if i['name'] != name]
            if not sched[day]:
                del sched[day]
            self.save_projects()
            self._update_manual_list()
            self._update_proj_list()
            self.refresh_daily_plan(auto=True)

    # ===== 文档导入 =====
    def import_plan_for_project(self):
        if not self.selected_project:
            messagebox.showwarning('提示', '请先选择项目！')
            return
        fp = filedialog.askopenfilename(
            title=f'为「{self.selected_project}」选择计划文件',
            filetypes=[('支持的文件', '*.txt;*.docx'), ('文本文件', '*.txt'),
                       ('Word文档', '*.docx'), ('所有文件', '*.*')])
        if not fp:
            return
        try:
            d = self._proj_dir(self.selected_project)
            ts = datetime.now().strftime("plan_%Y%m%d_%H%M%S")
            ext = os.path.splitext(fp)[1]
            dest = os.path.join(d, f"{ts}{ext}")
            shutil.copy2(fp, dest)
            self.projects[self.selected_project]['plans'].append(dest)
            self.save_projects()
            self._update_proj_list()
            self._update_doc_list()
            self.refresh_daily_plan(auto=True)
            # 解析校验：总天数、Day 是否连贯
            default_min = self.projects[self.selected_project].get(
                'task_minutes', DEFAULT_TASK_MINUTES)
            days = self._parse_doc_all(dest, default_min,
                                       self._proj_ref_start(self.selected_project))
            if not days:
                messagebox.showwarning('导入成功但未识别到计划',
                    '文件已保存，但没有识别出任何 Day 内容！\n'
                    '支持的格式：\n'
                    '① [ ] Day 1：任务1｜任务2｜任务3\n'
                    '② Day1 单独一行，后续每行：任务名,分钟数')
                return
            day_nums = sorted(days.keys())
            total_tasks = sum(len(v) for v in days.values())
            if getattr(self, '_last_parse_datebased', False):
                msg = (f'识别为「按日期排课」计划，共 {total_tasks} 次课程/安排！📅\n'
                       f'已按项目起始日期换算为 Day 编号（第{day_nums[0]}天 ~ 第{day_nums[-1]}天），'
                       f'有安排的当天会自动出现在首页打卡任务中。\n'
                       f'历史学习进度继续累加。')
            else:
                missing = [d for d in range(day_nums[0], day_nums[-1] + 1)
                           if d not in days]
                msg = (f'成功解析 {len(day_nums)} 天计划（Day{day_nums[0]} ~ Day{day_nums[-1]}），'
                       f'共 {total_tasks} 个分点任务！📄\n'
                       f'已更新为该项目的最新计划版本，历史学习进度继续累加。\n'
                       f'今天是第 {self._calc_project_day(self.selected_project)} 天，'
                       f'首页已显示新计划当天的内容。')
                if missing:
                    msg += (f'\n\n⚠️ 注意：Day 编号不连贯，缺少：'
                            f'{"、".join(map(str, missing[:20]))}'
                            f'{"……" if len(missing) > 20 else ""}')
            messagebox.showinfo('导入成功', msg)
        except Exception as e:
            messagebox.showerror('导入失败', str(e))

    def delete_project_doc(self):
        if not self.selected_project:
            return
        plans = self._plan_files(self.selected_project)
        if not plans:
            messagebox.showinfo('提示', '该项目没有可删除的文件')
            return
        win = tk.Toplevel(self.root)
        win.title(f'{self.selected_project} - 删除文档')
        win.geometry('500x380')
        win.configure(bg=PASTEL['bg'])
        win.transient(self.root)
        win.grab_set()
        tk.Label(win, text='选择要删除的文档', font=_bold(15), fg=PASTEL['title'],
                 bg=PASTEL['bg']).pack(pady=15)
        lb = tk.Listbox(win, font=_norm(12), bg=PASTEL['card_bg'], fg=PASTEL['text'],
            selectbackground=PASTEL['accent2'], relief='solid', bd=1)
        lb.pack(fill=tk.BOTH, expand=True, padx=25, pady=5)
        for p in plans:
            lb.insert(tk.END, os.path.basename(p))

        def do_del():
            s = lb.curselection()
            if not s:
                messagebox.showwarning('提示', '请选择要删除的文档！')
                return
            if messagebox.askyesno('确认删除', f'确定要删除 {os.path.basename(plans[s[0]])} 吗？'):
                os.remove(plans[s[0]])
                win.destroy()
                self._update_proj_list()
                self._update_doc_list()
                self.refresh_daily_plan(auto=True)

        ttk.Button(win, text='确认删除', command=do_del,
                   style='Delete.TButton').pack(side=tk.LEFT, padx=30, pady=15)
        ttk.Button(win, text='取消', command=win.destroy,
                   style='Cute.TButton').pack(side=tk.RIGHT, padx=30, pady=15)

    # ===== 文档解析 =====
    # 行内格式：[ ] Day 1：任务1｜任务2（兼容 markdown：- [ ] **Day 1**：…）
    # 支持 Day 区间（Day 1-2 / Day 15-21）与编号后缀（Day 7（周日复盘））
    INLINE_DAY_RE = re.compile(
        r'^[\s\-\*•>]*(?:\[\s*[xX✔]?\s*\]\s*)?\**'
        r'(?:Day|day|DAY|D)\s*(\d+)'
        r'(?:\s*[-–—~]\s*(\d+))?'
        r'\**[^：:]*[：:]\s*(.+)$')
    INLINE_DAY_CN_RE = re.compile(
        r'^[\s\-\*•>]*(?:\[\s*[xX✔]?\s*\]\s*)?\**'
        r'第\s*(\d+)\s*天'
        r'(?:\s*[-–—~]\s*(?:第\s*)?(\d+)\s*天)?'
        r'\**[^：:]*[：:]\s*(.+)$')
    # 块格式（旧格式）：Day1 / 第1天 单独一行，后续行是"任务,分钟"
    BLOCK_DAY_RE = re.compile(
        r'^(?:Day\s*(\d+)|第\s*(\d+)\s*天|D\s*(\d+))\s*$', re.IGNORECASE)
    # 日期排课格式：7.25 瑜伽 16:00-17:30
    DATE_LINE_RE = re.compile(
        r'^(\d{1,2})\s*[\.、/]\s*(\d{1,2})\s+'
        r'(.+?)\s+(\d{1,2})[:：](\d{2})\s*[-–—~]\s*(\d{1,2})[:：](\d{2})\s*$')

    def _read_doc_lines(self, fp):
        """读取 txt/docx 全部非空行"""
        all_lines = []
        try:
            if fp.endswith('.docx'):
                if not DOCX_SUPPORT:
                    return []
                doc = Document(fp)
                for para in doc.paragraphs:
                    l = para.text.strip()
                    if l:
                        all_lines.append(l)
                for t in doc.tables:
                    for row in t.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if len(cells) >= 2:
                            all_lines.append(f"{cells[0]},{cells[1]}")
                        elif len(cells) == 1:
                            all_lines.append(cells[0])
            else:
                content = ''
                for enc in ('utf-8', 'utf-8-sig', 'gbk'):
                    try:
                        with open(fp, 'r', encoding=enc) as f:
                            content = f.read()
                        if content.strip():
                            break
                    except Exception:
                        continue
                for l in content.split('\n'):
                    l = l.strip()
                    if l:
                        all_lines.append(l)
        except Exception:
            return []
        return all_lines

    def _extract_task_minutes(self, text, default_min):
        """从任务文本提取时长：优先「限时X分钟」，其次「X小时」，否则用默认值"""
        m = re.search(r'限时\s*(\d+)\s*分钟', text)
        if m:
            return int(m.group(1))
        m = re.search(r'(\d+(?:\.\d+)?)\s*小时', text)
        if m:
            return int(round(float(m.group(1)) * 60))
        m = re.search(r'(?:约|大概|预计)\s*(\d+)\s*分钟', text)
        if m:
            return int(m.group(1))
        return default_min

    def _split_subtasks(self, content):
        """把 Day 行内容按 ｜/|/；/; 拆成分点任务，并去掉 markdown 加粗等标记"""
        parts = re.split(r'[｜|；;]+', content)
        out = []
        for p in parts:
            p = p.strip().strip('。').replace('**', '').replace('__', '').strip()
            if p:
                out.append(p)
        return out

    def _parse_doc_all(self, fp, default_min=DEFAULT_TASK_MINUTES, ref_start=None):
        """
        解析整个计划文档 → {天数(int): [{'name':..., 'time':...}, ...]}
        支持三种格式：
          1. 行内格式：[ ] Day 1：任务1｜任务2（兼容 - [ ] **Day 1** markdown，
             支持 Day 区间 Day 1-2、编号后缀 Day 7（复盘））
          2. 块格式：Day1 单独一行，后续每行「任务名,分钟数」
          3. 日期排课格式：7.25 瑜伽 16:00-17:30（需传 ref_start=项目起始日期，
             自动换算 Day 编号，时长按时间段计算）
        策略说明、周标题、表格等非 Day 行会被自动忽略。
        """
        all_lines = self._read_doc_lines(fp)
        days = {}
        block_mode = False
        current_day = None
        self._last_parse_datebased = False

        for line in all_lines:
            # ③ 日期排课格式
            md = self.DATE_LINE_RE.match(line)
            if md and ref_start:
                self._last_parse_datebased = True
                mon, day = int(md.group(1)), int(md.group(2))
                name = md.group(3).strip().replace('**', '')
                h1, m1, h2, m2 = (int(md.group(i)) for i in range(4, 8))
                mins = (h2 * 60 + m2) - (h1 * 60 + m1)
                try:
                    entry_date = date(ref_start.year, mon, day)
                    if entry_date < ref_start:
                        entry_date = date(ref_start.year + 1, mon, day)
                except ValueError:
                    continue
                dnum = (entry_date - ref_start).days + 1
                if dnum >= 1:
                    days.setdefault(dnum, []).append(
                        {'name': name, 'time': mins if mins > 0 else default_min})
                continue

            # ① 行内格式（含 Day 区间）
            m = self.INLINE_DAY_RE.match(line) or self.INLINE_DAY_CN_RE.match(line)
            if m:
                d_start = int(m.group(1))
                d_end = int(m.group(2)) if m.group(2) else d_start
                block_mode = False
                current_day = d_start
                for d in range(d_start, d_end + 1):
                    tasks = days.setdefault(d, [])
                    for sub in self._split_subtasks(m.group(3)):
                        tasks.append({'name': sub,
                                      'time': self._extract_task_minutes(sub, default_min)})
                continue

            # ② 块格式
            m = self.BLOCK_DAY_RE.match(line)
            if m:
                current_day = int(m.group(1) or m.group(2) or m.group(3))
                block_mode = True
                days.setdefault(current_day, [])
                continue
            if block_mode and current_day is not None:
                name, time_str = self._parse_line(line)
                if name and time_str and time_str.isdigit():
                    days[current_day].append({'name': name, 'time': int(time_str)})
            # 非 Day 行（策略/周标题/表格等）在行内模式下直接忽略

        return days

    def _parse_doc_by_day(self, fp, day_num, default_min=DEFAULT_TASK_MINUTES,
                          ref_start=None):
        """解析文档中指定Day的内容 → {任务名: {'time': 分钟}}（同一天同名任务合并）"""
        days = self._parse_doc_all(fp, default_min, ref_start)
        parsed = {}
        for idx, item in enumerate(days.get(day_num, []), 1):
            name = item['name'] if item['name'] not in parsed \
                else f"{item['name']} ({idx})"
            parsed[name] = {'time': item['time']}
        return parsed

    def _parse_line(self, line):
        line = line.strip()
        if not line:
            return None, None
        if ',' in line:
            parts = [p.strip() for p in line.split(',', 1)]
        elif '，' in line:
            parts = [p.strip() for p in line.split('，', 1)]
        elif '\t' in line:
            parts = [p.strip() for p in line.split('\t', 1)]
        elif '  ' in line:
            parts = [p.strip() for p in re.split(r'\s{2,}', line, 1)]
        elif ' ' in line:
            idx = line.rfind(' ')
            parts = [line[:idx].strip(), line[idx:].strip()]
        else:
            return line, None
        if len(parts) >= 2:
            name = parts[0].strip()
            d = re.findall(r'\d+', parts[1].strip())
            return name, (d[0] if d else None)
        elif len(parts) == 1 and parts[0]:
            return parts[0], None
        return None, None

    # ===== 智能进度结转 =====
    def _project_past_totals(self, proj_name):
        """统计某项目在今天之前的：计划总分钟 与 实际总分钟（按复盘记录）"""
        planned = 0
        actual = 0
        prefix = f"{proj_name} - "
        for fp in glob.glob(os.path.join(self.daily_dir, "*.json")):
            day_str = os.path.splitext(os.path.basename(fp))[0]
            if day_str >= self.today:
                continue
            try:
                with open(fp, 'r', encoding='utf-8') as f:
                    d = json.load(f)
            except Exception:
                continue
            for key, info in d.get('plan', {}).items():
                if info.get('project') == proj_name or key.startswith(prefix):
                    planned += info.get('time', 0)
            for key, recs in d.get('review', {}).items():
                if key.startswith(prefix):
                    actual += sum(r.get('time', 0) for r in recs)
        return planned, actual

    def _compute_carryover(self, proj_name):
        """结转分钟数：>0 表示落后（今天加量），<0 表示超前（今天减量）"""
        planned, actual = self._project_past_totals(proj_name)
        return planned - actual

    def _proj_ref_start(self, proj_name):
        """项目起始日期（date对象），用于日期排课格式换算Day编号"""
        sd = self.projects.get(proj_name, {}).get('start_date')
        try:
            return datetime.strptime(sd, "%Y-%m-%d").date()
        except Exception:
            return date.today()

    def _today_base_schedule(self, proj_name):
        """合并「文档计划 + 手动计划」得到今天（第N天）原计划：{任务名: 分钟}"""
        day_num = self._calc_project_day(proj_name)
        default_min = self.projects[proj_name].get('task_minutes', DEFAULT_TASK_MINUTES)
        items = {}
        plans = self._plan_files(proj_name)
        if plans:
            try:
                for name, info in self._parse_doc_by_day(
                        plans[0], day_num, default_min,
                        self._proj_ref_start(proj_name)).items():
                    items[name] = items.get(name, 0) + info['time']
            except Exception:
                pass
        sched = self.projects[proj_name].get('manual_schedule', {})
        for it in sched.get(str(day_num), []):
            items[it['name']] = items.get(it['name'], 0) + it['time']
        return items

    def refresh_daily_plan(self, auto=False):
        """
        只生成/同步「今天」的任务：
        - 每个项目取今天（第N天）的原计划
        - 按以往「计划-实际」差额自动增减（超前则减量、落后则补量）
        - 已存在的任务保留完成状态，未被安排的旧条目若未完成则移除
        """
        self.adjust_notes = []
        scheduled_keys = set()
        changed = False

        for proj_name in sorted(self.projects.keys()):
            day_num = self._calc_project_day(proj_name)
            base_items = self._today_base_schedule(proj_name)
            carry = self._compute_carryover(proj_name)

            base_total = sum(base_items.values())
            target_total = max(0, base_total + carry)

            if base_total > 0:
                # 把差额按比例分摊到今天的各项任务
                delta = target_total - base_total
                names = list(base_items.keys())
                allocated, acc = {}, 0
                for i, nm in enumerate(names):
                    if i < len(names) - 1:
                        a = int(round(delta * base_items[nm] / base_total))
                        acc += a
                    else:
                        a = delta - acc
                    allocated[nm] = max(0, base_items[nm] + a)

                if carry < 0:
                    self.adjust_notes.append(
                        f"🎉「{proj_name}」之前多学了 {-carry} 分钟，今日计划减少 {-carry} 分钟")
                elif carry > 0:
                    self.adjust_notes.append(
                        f"💪「{proj_name}」之前还差 {carry} 分钟，今日计划已补上")

                for nm, mins in allocated.items():
                    key = f"{proj_name} - {nm}"
                    scheduled_keys.add(key)
                    base = base_items[nm]
                    adj = mins - base
                    if key in self.daily_plan:
                        info = self.daily_plan[key]
                        if not info.get('completed') and (info.get('time') != mins
                                                          or info.get('base') != base):
                            info.update(time=mins, base=base, adjust=adj, day=day_num)
                            changed = True
                    else:
                        self.daily_plan[key] = {
                            'project': proj_name, 'task': nm,
                            'time': mins, 'base': base, 'adjust': adj,
                            'completed': False, 'day': day_num, 'source': 'schedule',
                        }
                        changed = True
            elif carry > 0:
                # 今天原本无安排但有历史欠账 → 生成补学任务
                key = f"{proj_name} - 补学任务"
                scheduled_keys.add(key)
                self.adjust_notes.append(
                    f"💪「{proj_name}」之前还差 {carry} 分钟，今日安排了补学任务")
                if key not in self.daily_plan:
                    self.daily_plan[key] = {
                        'project': proj_name, 'task': '补学任务',
                        'time': carry, 'base': 0, 'adjust': carry,
                        'completed': False, 'day': day_num, 'source': 'schedule',
                    }
                    changed = True

        # 移除今天不再属于计划安排的未完成条目（手动临时任务保留）
        for key in list(self.daily_plan.keys()):
            info = self.daily_plan[key]
            if info.get('source') == 'schedule' and key not in scheduled_keys \
                    and not info.get('completed'):
                del self.daily_plan[key]
                changed = True

        if changed or not auto:
            self.save_daily_data()
        self._update_checkin_tree()
        self._update_review_combo()
        self._update_stat_cards()

        if self.adjust_notes:
            self.adjust_var.set('✨ 智能调整：' + '　'.join(self.adjust_notes))
        else:
            self.adjust_var.set('✨ 今日计划与原定一致，保持稳定节奏哦~'
                                if self.daily_plan else '🌱 还没有计划，去「学习计划」页添加项目吧~')

        if not auto:
            note = '\n'.join(self.adjust_notes) if self.adjust_notes else '今日计划与原定一致。'
            messagebox.showinfo('同步完成', f'今日计划已同步！\n{note}')

    # ===== 数据持久化 =====
    def load_daily_data(self):
        df = os.path.join(self.daily_dir, f"{self.today}.json")
        if os.path.exists(df):
            try:
                with open(df, 'r', encoding='utf-8') as f:
                    d = json.load(f)
                    self.daily_plan = d.get('plan', {})
                    self.review_data = d.get('review', {})
            except Exception:
                self.daily_plan = {}
                self.review_data = {}
        self._load_historical()
        self._update_checkin_tree()
        self._update_review_tree()
        self._update_review_combo()
        self._update_stat_cards()
        self.update_review_chart()

    def save_daily_data(self):
        df = os.path.join(self.daily_dir, f"{self.today}.json")
        with open(df, 'w', encoding='utf-8') as f:
            json.dump({'plan': self.daily_plan, 'review': self.review_data},
                      f, ensure_ascii=False, indent=2)

    def _load_historical(self):
        self.all_review_data = {}
        today_file = os.path.join(self.daily_dir, f"{self.today}.json")
        for fp in glob.glob(os.path.join(self.daily_dir, "*.json")):
            if fp == today_file:
                continue
            try:
                with open(fp, 'r', encoding='utf-8') as f:
                    d = json.load(f)
                    for k, rs in d.get('review', {}).items():
                        self.all_review_data.setdefault(k, 0)
                        for r in rs:
                            self.all_review_data[k] += r.get('time', 0)
            except Exception:
                continue

    # ===== 今日任务操作 =====
    def add_daily_item(self):
        name = self.checkin_name_var.get().strip()
        time_str = self.checkin_time_var.get().strip()
        if not name:
            messagebox.showwarning('提示', '请输入任务名称！')
            return
        if not time_str.isdigit():
            messagebox.showwarning('提示', '请输入有效的时间（数字分钟）！')
            return
        if name in self.daily_plan:
            messagebox.showwarning('提示', '该任务已存在！')
            return
        t = int(time_str)
        self.daily_plan[name] = {
            'project': None, 'task': name, 'time': t, 'base': t, 'adjust': 0,
            'completed': False, 'source': 'user',
        }
        self.save_daily_data()
        self._update_checkin_tree()
        self._update_review_combo()
        self._update_stat_cards()
        self.checkin_name_var.set('')
        self.checkin_time_var.set('')

    def delete_daily_item(self):
        sel = self.checkin_tree.selection()
        if not sel:
            messagebox.showwarning('提示', '请先选择要删除的任务！')
            return
        key = self.checkin_tree.item(sel[0], 'values')[0]
        if messagebox.askyesno('删除确认', f'确定要删除「{key}」吗？'):
            if key in self.daily_plan:
                if key == self.timer_task_key:
                    self.timer_running = False
                    self._reset_timer_state()
                del self.daily_plan[key]
                self.save_daily_data()
                self._update_checkin_tree()
                self._update_review_combo()
                self._update_stat_cards()

    def _update_checkin_tree(self):
        for item in self.checkin_tree.get_children():
            self.checkin_tree.delete(item)
        for key, info in self.daily_plan.items():
            base = info.get('base', info['time'])
            adj = info.get('adjust', 0)
            if adj > 0:
                adj_str = f'+{adj} 分钟'
            elif adj < 0:
                adj_str = f'{adj} 分钟'
            else:
                adj_str = '—'
            is_current = (key == self.timer_task_key)
            if info['completed']:
                status, timer_col, action = '✅ 已完成', '—', '已打卡'
                tags = ('done',)
            else:
                status = '🌱 进行中'
                action = '👉 打卡'
                if is_current and self.timer_running:
                    timer_col, tags = '⏸ 暂停', ('timing',)
                elif is_current:
                    timer_col, tags = '▶ 继续', ('timing',)
                else:
                    timer_col, tags = '▶ 开始', ()
            self.checkin_tree.insert('', tk.END, values=(
                key, f'{base} 分钟', adj_str, f"{info['time']} 分钟",
                status, timer_col, action), tags=tags)
        self._update_summary_tree()

    def _on_checkin_click(self, event):
        region = self.checkin_tree.identify_region(event.x, event.y)
        if region != 'cell':
            return
        col = self.checkin_tree.identify_column(event.x)
        item = self.checkin_tree.identify_row(event.y)
        if not item:
            return
        vals = self.checkin_tree.item(item, 'values')
        key = vals[0]
        if col == '#6' and '进行中' in str(vals[4]):
            self._on_row_timer_click(key)          # 「计时」列：开始/暂停/切换
        elif col == '#7' and '进行中' in str(vals[4]):
            self.checkin_item(key)                 # 「操作」列：打卡

    def _on_row_timer_click(self, key):
        """行内计时按钮：同一任务=暂停/继续；不同任务=自动停止当前并切换"""
        if key not in self.daily_plan or self.daily_plan[key].get('completed'):
            return
        if key == self.timer_task_key:
            if self.timer_running:
                self.pause_timer()
            else:
                self._resume_timer()
            return
        # 交叉完成：先停掉正在计时的任务，自动记录已学时长
        if self.timer_task_key:
            self._auto_stop_current(reason='切换任务')
        self._start_timing(key)

    def checkin_item(self, key, silent=False):
        if key in self.daily_plan and not self.daily_plan[key]['completed']:
            # 若该任务正在计时，先自动停止并记录已学时长
            if key == self.timer_task_key:
                self._auto_stop_current(reason='打卡')
                self._reset_timer_state()
            self.daily_plan[key]['completed'] = True
            self.save_daily_data()
            self._update_checkin_tree()
            self._update_stat_cards()
            if not silent:
                messagebox.showinfo('打卡成功',
                    f'「{key}」打卡完成！🎉\n{random.choice(ENCOURAGE)}')

    def _update_stat_cards(self):
        total = len(self.daily_plan)
        done = sum(1 for i in self.daily_plan.values() if i.get('completed'))
        mins = sum(i.get('time', 0) for i in self.daily_plan.values())
        self.stat_cards['total'].config(text=str(total))
        self.stat_cards['done'].config(text=str(done))
        self.stat_cards['mins'].config(text=f'{mins}分')
        self.stat_cards['streak'].config(text=f'{self._calc_streak()}天')

    def _calc_streak(self):
        """连续学习天数：从今天(或昨天)起向前，凡当天有打卡或有复盘记录即算"""
        streak = 0
        d = date.today()
        # 今天还没学不破坏连续记录，从昨天开始也算
        if not any(i.get('completed') for i in self.daily_plan.values()) \
                and not self.review_data:
            d -= timedelta(days=1)
        while True:
            fp = os.path.join(self.daily_dir, d.strftime("%Y-%m-%d") + ".json")
            if not os.path.exists(fp):
                break
            try:
                with open(fp, 'r', encoding='utf-8') as f:
                    data = json.load(f)
            except Exception:
                break
            has_done = any(i.get('completed') for i in data.get('plan', {}).values())
            has_review = bool(data.get('review'))
            if has_done or has_review:
                streak += 1
                d -= timedelta(days=1)
            else:
                break
        return streak

    # ===== 番茄钟 =====
    def _start_timing(self, key):
        """开始为指定任务倒计时"""
        self.timer_task_key = key
        self.timer_remaining = self.daily_plan[key]['time'] * 60
        self.timer_total = self.timer_remaining
        self.timer_running = True
        self.pause_timer_btn.config(state=tk.NORMAL)
        self.stop_timer_btn.config(state=tk.NORMAL)
        self.checkin_status_var.set(f'🍅 专注中：{key}')
        self._update_checkin_tree()
        self._schedule_tick()

    def _resume_timer(self):
        if not self.timer_task_key or self.timer_remaining <= 0:
            return
        self.timer_running = True
        self.pause_timer_btn.config(state=tk.NORMAL)
        self.stop_timer_btn.config(state=tk.NORMAL)
        self.checkin_status_var.set(f'🍅 专注中：{self.timer_task_key}')
        self._update_checkin_tree()
        self._schedule_tick()

    def _schedule_tick(self):
        """主线程驱动的秒级倒计时（root.after，线程安全）"""
        self._cancel_tick()
        self._timer_after_id = self.root.after(1000, self._tick)

    def _cancel_tick(self):
        if self._timer_after_id is not None:
            try:
                self.root.after_cancel(self._timer_after_id)
            except Exception:
                pass
            self._timer_after_id = None

    def _tick(self):
        self._timer_after_id = None
        if not self.timer_running:
            return
        self.timer_remaining -= 1
        if self.timer_remaining <= 0:
            self.timer_remaining = 0
            self.timer_running = False
            self.timer_display.config(text='00:00:00')
            self._timer_done()
            return
        h = self.timer_remaining // 3600
        m = (self.timer_remaining % 3600) // 60
        s = self.timer_remaining % 60
        self.timer_display.config(text=f"{h:02d}:{m:02d}:{s:02d}")
        self._timer_after_id = self.root.after(1000, self._tick)

    def _auto_stop_current(self, reason='切换任务'):
        """自动停止当前计时任务，已专注时长（≥1分钟）自动记入复盘"""
        self.timer_running = False
        self._cancel_tick()
        key = self.timer_task_key
        elapsed = max(0, self.timer_total - self.timer_remaining)
        mins = elapsed // 60
        self.timer_remaining = 0
        self.timer_total = 0
        if key and mins > 0:
            self._record_review(key, mins, f'番茄钟{reason}')
            self.checkin_status_var.set(f'📝 已自动记录「{key}」{mins} 分钟')
        return key, mins

    def _reset_timer_state(self):
        self.timer_task_key = None
        self.timer_remaining = 0
        self.timer_total = 0
        self.timer_display.config(text='00:00:00')
        self.pause_timer_btn.config(state=tk.DISABLED)
        self.stop_timer_btn.config(state=tk.DISABLED)
        self._update_checkin_tree()

    def _timer_done(self):
        key = self.timer_task_key
        mins = self.timer_total // 60
        self._reset_timer_state()
        if key and key in self.daily_plan:
            self._record_review(key, mins, '番茄钟完成')
            self.checkin_item(key, silent=True)
            self.checkin_status_var.set(random.choice(ENCOURAGE))
            messagebox.showinfo('🎉 太棒了！',
                f'「{key}」专注完成 {mins} 分钟，已自动打卡并记录复盘！\n{random.choice(ENCOURAGE)}')

    def pause_timer(self):
        if not self.timer_running:
            return
        self.timer_running = False
        self._cancel_tick()
        self.pause_timer_btn.config(state=tk.DISABLED)
        self.checkin_status_var.set('☕ 休息一下，点击「▶ 继续」回来哦~')
        self._update_checkin_tree()

    def stop_timer(self):
        """手动结束：把已专注的时间记入复盘"""
        if not self.timer_task_key:
            return
        key = self.timer_task_key
        elapsed = max(0, self.timer_total - self.timer_remaining)
        mins = elapsed // 60
        self.timer_running = False
        self._reset_timer_state()
        if mins > 0:
            if messagebox.askyesno('结束计时',
                    f'本次已专注 {mins} 分钟，要记入复盘吗？\n（任务「{key}」）'):
                self._record_review(key, mins, '番茄钟提前结束')
                messagebox.showinfo('已记录', f'已记录 {mins} 分钟实际学习！🌟')
        self.checkin_status_var.set(random.choice(ENCOURAGE))

    def _record_review(self, key, mins, content):
        self.review_data.setdefault(key, []).append({
            'time': int(mins), 'content': content,
            'timestamp': datetime.now().strftime("%H:%M:%S"),
        })
        self.save_daily_data()
        self._load_historical()
        self._update_review_tree()
        self._update_stat_cards()
        self.update_review_chart()

    # ===== 复盘（支持 添加/修改/删除） =====
    def _update_review_combo(self):
        self.review_name_combo['values'] = list(self.daily_plan.keys())

    def add_review_item(self):
        name = self.review_name_var.get().strip()
        ts = self.review_time_var.get().strip()
        ct = self.review_content_var.get().strip()
        if not name:
            messagebox.showwarning('提示', '请选择学习项目！')
            return
        if not ts.isdigit():
            messagebox.showwarning('提示', '请输入有效的时间（数字）！')
            return
        if self._editing_review is not None:
            # 修改模式：更新原记录（保留原时间戳）
            old_name, idx = self._editing_review
            recs = self.review_data.get(old_name, [])
            if 0 <= idx < len(recs):
                old_ts = recs[idx].get('timestamp',
                                       datetime.now().strftime("%H:%M:%S"))
                del recs[idx]
                if not recs:
                    del self.review_data[old_name]
                self.review_data.setdefault(name, []).append({
                    'time': int(ts), 'content': ct, 'timestamp': old_ts})
            self._exit_edit_mode()
            msg = '复盘记录已修改！✏️'
        else:
            self.review_data.setdefault(name, []).append({
                'time': int(ts), 'content': ct,
                'timestamp': datetime.now().strftime("%H:%M:%S"),
            })
            msg = f'复盘记录已保存！{random.choice(ENCOURAGE)}'
        self.save_daily_data()
        self._load_historical()
        self._update_review_tree()
        self._update_stat_cards()
        self.update_review_chart()
        self.review_time_var.set('')
        self.review_content_var.set('')
        messagebox.showinfo('成功', msg)

    def edit_review_item(self):
        sel = self.review_tree.selection()
        if not sel:
            messagebox.showwarning('提示', '请先选中要修改的复盘记录！')
            return
        name, idx = sel[0].rsplit('||', 1)
        recs = self.review_data.get(name, [])
        if not recs or int(idx) >= len(recs):
            return
        r = recs[int(idx)]
        self.review_name_var.set(name)
        self.review_time_var.set(str(r['time']))
        self.review_content_var.set(r.get('content', ''))
        self._editing_review = (name, int(idx))
        self.add_review_btn.config(text='💾 保存修改')

    def _exit_edit_mode(self):
        self._editing_review = None
        self.add_review_btn.config(text='添加复盘')

    def delete_review_item(self):
        sel = self.review_tree.selection()
        if not sel:
            messagebox.showwarning('提示', '请先选中要删除的复盘记录！')
            return
        name, idx = sel[0].rsplit('||', 1)
        recs = self.review_data.get(name, [])
        if not recs or int(idx) >= len(recs):
            return
        r = recs[int(idx)]
        if messagebox.askyesno('删除确认',
                f'确定删除这条复盘吗？\n「{name}」{r["time"]}分钟 {r.get("content","")}'):
            del recs[int(idx)]
            if not recs:
                del self.review_data[name]
            self._exit_edit_mode()
            self.save_daily_data()
            self._load_historical()
            self._update_review_tree()
            self._update_stat_cards()
            self.update_review_chart()

    def _update_review_tree(self):
        for item in self.review_tree.get_children():
            self.review_tree.delete(item)
        for name, recs in self.review_data.items():
            for i, r in enumerate(recs):
                ts = r.get('timestamp', '')
                content = f"[{ts}] {r['content']}" if ts else r['content']
                self.review_tree.insert('', tk.END, iid=f'{name}||{i}',
                    values=(name, f"{r['time']} 分钟", content))
        self._update_summary_tree()

    def _update_summary_tree(self):
        """按项目汇总今日：计划总时长 vs 实际总时长，差额将自动结转到之后天数"""
        for item in self.summary_tree.get_children():
            self.summary_tree.delete(item)
        groups = {}   # 项目 -> {'planned': x, 'actual': y}
        for key, info in self.daily_plan.items():
            proj = info.get('project') or '临时任务'
            groups.setdefault(proj, {'planned': 0, 'actual': 0})
            groups[proj]['planned'] += info.get('time', 0)
        for key, recs in self.review_data.items():
            info = self.daily_plan.get(key)
            if info:
                proj = info.get('project') or '临时任务'
            else:
                proj = key.split(' - ')[0] if ' - ' in key else '临时任务'
                if proj not in self.projects:
                    proj = '临时任务'
            groups.setdefault(proj, {'planned': 0, 'actual': 0})
            groups[proj]['actual'] += sum(r.get('time', 0) for r in recs)
        for proj, g in sorted(groups.items()):
            diff = g['planned'] - g['actual']   # >0 少学了，<0 多学了
            if diff > 0:
                diff_str, effect = f'+{diff} 分钟', f'📌 少学了 {diff} 分钟，明天计划将增加 {diff} 分钟'
            elif diff < 0:
                diff_str, effect = f'{diff} 分钟', f'🎉 多学了 {-diff} 分钟，明天计划将减少 {-diff} 分钟'
            else:
                diff_str, effect = '0', '✨ 按计划完成，后续计划不变'
            if proj == '临时任务':
                effect = '不参与智能结转'
            self.summary_tree.insert('', tk.END, values=(
                proj, f"{g['planned']} 分钟", f"{g['actual']} 分钟", diff_str, effect))

    def _short_name(self, name, n=14):
        return name if len(name) <= n else name[:n - 1] + '…'

    def _project_total_minutes(self, proj_name):
        """项目计划总时长 = 最新计划文档全部Day + 手动计划全部Day（带缓存）"""
        total = 0
        plans = self._plan_files(proj_name)
        if plans:
            fp = plans[0]
            try:
                mtime = os.path.getmtime(fp)
            except Exception:
                mtime = 0
            ref = self._proj_ref_start(proj_name)
            cache_key = (fp, mtime, self.projects[proj_name].get(
                'task_minutes', DEFAULT_TASK_MINUTES), str(ref))
            if not hasattr(self, '_doc_cache'):
                self._doc_cache = {}
            if cache_key not in self._doc_cache:
                days = self._parse_doc_all(fp, cache_key[2], ref)
                self._doc_cache.clear()
                self._doc_cache[cache_key] = sum(
                    t['time'] for tasks in days.values() for t in tasks)
            total += self._doc_cache[cache_key]
        for tasks in self.projects[proj_name].get('manual_schedule', {}).values():
            total += sum(t['time'] for t in tasks)
        return total

    def _project_done_minutes(self, proj_name):
        """项目累计实际学习时长（历史 + 今天复盘）"""
        prefix = f"{proj_name} - "
        done = sum(v for k, v in self.all_review_data.items() if k.startswith(prefix))
        for k, recs in self.review_data.items():
            if k.startswith(prefix):
                done += sum(r.get('time', 0) for r in recs)
        return done

    def update_review_chart(self):
        self.ax.clear()
        self.ax.set_facecolor(PASTEL['chart_bg'])
        mode = self.chart_mode_var.get()
        colors = PASTEL['chart_colors']
        if mode == 'today':
            self._draw_today_24h(colors)
        else:
            self._draw_overall_progress(colors)
        self.figure.tight_layout()
        self.canvas.draw()

    def _draw_today_24h(self, colors):
        """今日24小时分布：横轴0-24时，每个任务一行，按打卡时间戳画出学习时段"""
        if not self.review_data:
            self.ax.text(0.5, 0.5, '暂无数据，去学习一会吧~ 🌱', ha='center', va='center',
                fontproperties=CHART_FONT_TITLE, color=PASTEL['accent'])
            self.figure.suptitle('今日24小时学习分布 ⏰',
                fontproperties=CHART_FONT_TITLE, color=PASTEL['title'])
            return
        tasks = list(self.review_data.keys())
        for i, name in enumerate(tasks):
            spans = []
            for r in self.review_data[name]:
                ts = r.get('timestamp', '00:00:00')
                try:
                    hh, mm = int(ts[:2]), int(ts[3:5])
                except Exception:
                    hh, mm = 0, 0
                start = hh + mm / 60.0
                dur = max(r.get('time', 0) / 60.0, 0.05)   # 太短也给个可见宽度
                spans.append((start, dur))
            self.ax.broken_barh(spans, (i - 0.36, 0.72),
                facecolors=colors[i % len(colors)], edgecolor='white', linewidth=1)
            total = sum(r.get('time', 0) for r in self.review_data[name])
            self.ax.text(24.3, i, f'{total}分钟', va='center', ha='left',
                fontproperties=CHART_FONT_SM, color=PASTEL['text'])
        self.ax.set_yticks(range(len(tasks)))
        self.ax.set_yticklabels([self._short_name(t) for t in tasks],
            fontproperties=CHART_FONT_SM, color=PASTEL['text'])
        self.ax.set_ylim(-0.6, len(tasks) - 0.4)
        self.ax.invert_yaxis()
        self.ax.set_xlim(0, 26)
        self.ax.set_xticks(range(0, 25, 2))
        self.ax.set_xticklabels([f'{h}时' for h in range(0, 25, 2)],
            fontproperties=CHART_FONT_SM, color=PASTEL['text'])
        now = datetime.now()
        self.ax.axvline(now.hour + now.minute / 60.0, color=PASTEL['title'],
            linestyle='--', linewidth=1.2, alpha=0.7)
        self.ax.grid(axis='x', color=PASTEL['border'], linewidth=0.6, alpha=0.6)
        for sp in ['top', 'right', 'left']:
            self.ax.spines[sp].set_visible(False)
        self.ax.spines['bottom'].set_color(PASTEL['border'])
        self.figure.suptitle('今日24小时学习分布 ⏰（虚线=当前时间）',
            fontproperties=CHART_FONT_TITLE, color=PASTEL['title'])

    def _draw_overall_progress(self, colors):
        """整体进度：各长期项目 已完成/计划总量 进度条（不含临时任务）"""
        projs = sorted(self.projects.keys())
        if not projs:
            self.ax.text(0.5, 0.5, '还没有学习项目，去「学习计划」页添加吧~ 🌱',
                ha='center', va='center',
                fontproperties=CHART_FONT_TITLE, color=PASTEL['accent'])
            self.figure.suptitle('整体学习进度 📈',
                fontproperties=CHART_FONT_TITLE, color=PASTEL['title'])
            return
        names, pcts, labels = [], [], []
        for p in projs:
            total = self._project_total_minutes(p)
            done = self._project_done_minutes(p)
            pct = min(done / total * 100, 100) if total > 0 else 0
            names.append(p)
            pcts.append(pct)
            labels.append(f'{pct:.1f}%（{done}/{total}分钟）' if total > 0
                          else f'{done}分钟（暂无计划文档）')
        ypos = range(len(names))
        # 背景条（100%）+ 完成条
        self.ax.barh(ypos, [100] * len(names), color=PASTEL['frame_bg'],
            edgecolor=PASTEL['border'], height=0.6)
        bars = self.ax.barh(ypos, pcts, color=colors[:len(names)],
            edgecolor='white', height=0.6)
        for y, lb in zip(ypos, labels):
            self.ax.text(101, y, lb, va='center', ha='left',
                fontproperties=CHART_FONT_SM, color=PASTEL['text'])
        self.ax.set_yticks(list(ypos))
        self.ax.set_yticklabels(names, fontproperties=CHART_FONT, color=PASTEL['text'])
        self.ax.invert_yaxis()
        self.ax.set_xlim(0, 145)
        self.ax.set_xticks([0, 25, 50, 75, 100])
        self.ax.set_xticklabels(['0%', '25%', '50%', '75%', '100%'],
            fontproperties=CHART_FONT_SM, color=PASTEL['text'])
        for sp in ['top', 'right', 'left']:
            self.ax.spines[sp].set_visible(False)
        self.ax.spines['bottom'].set_color(PASTEL['border'])
        self.figure.suptitle('整体学习进度 📈（实际学习时长 / 计划总时长）',
            fontproperties=CHART_FONT_TITLE, color=PASTEL['title'])

    # ===== 设置 =====
    def browse_data_dir(self):
        d = filedialog.askdirectory(title='选择数据文件夹')
        if d:
            self.data_dir_var.set(d)

    def set_default_dir(self):
        self.data_dir_var.set(os.path.join(
            os.path.dirname(os.path.abspath(__file__)), "example"))

    def save_settings_and_restart(self):
        nd = self.data_dir_var.get().strip()
        if not nd:
            messagebox.showwarning('提示', '请输入有效的文件夹路径！')
            return
        if nd != self.data_dir:
            self.data_dir = nd
            self.daily_dir = os.path.join(self.data_dir, "daily")
            self.projects_file = os.path.join(self.data_dir, "projects.json")
            self.settings_file = os.path.join(self.data_dir, "settings.json")
            self._ensure_dirs()
            self.save_settings()
            messagebox.showinfo('设置已保存', '设置已保存，请重新启动应用生效！')
        else:
            messagebox.showinfo('设置已保存', '设置已保存！')


if __name__ == "__main__":
    root = tk.Tk()
    app = StudyTrackerApp(root)
    root.mainloop()
